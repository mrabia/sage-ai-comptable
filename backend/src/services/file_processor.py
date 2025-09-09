"""
Service de traitement des fichiers attachés
Supporte Excel, CSV, PDF, Images, Word, et autres formats
"""

import os
import json
import pandas as pd
import mimetypes
from typing import Dict, Any, List, Optional, Tuple
from werkzeug.utils import secure_filename
from datetime import datetime
import logging

# Import pour différents types de fichiers
try:
    from PIL import Image
    import pytesseract  # OCR pour images
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

logger = logging.getLogger(__name__)

class FileProcessorService:
    """Service de traitement et d'analyse des fichiers attachés"""
    
    def __init__(self, upload_folder: str = 'uploads'):
        self.upload_folder = upload_folder
        self.max_file_size = 50 * 1024 * 1024  # 50MB max
        self.allowed_extensions = {
            # Tableurs
            '.xlsx', '.xls', '.csv',
            # Documents
            '.pdf', '.docx', '.doc', '.txt',
            # Images
            '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff',
            # Autres formats de données
            '.json', '.xml'
        }
        
        # Créer le dossier d'upload s'il n'existe pas
        os.makedirs(self.upload_folder, exist_ok=True)
    
    def is_allowed_file(self, filename: str) -> bool:
        """Vérifier si le fichier est autorisé"""
        return any(filename.lower().endswith(ext) for ext in self.allowed_extensions)
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Obtenir les informations de base sur un fichier"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Fichier non trouvé: {file_path}")
        
        stat = os.stat(file_path)
        filename = os.path.basename(file_path)
        file_extension = os.path.splitext(filename)[1].lower()
        mime_type, _ = mimetypes.guess_type(file_path)
        
        return {
            'filename': filename,
            'file_size': stat.st_size,
            'file_extension': file_extension,
            'mime_type': mime_type or 'application/octet-stream',
            'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat()
        }
    
    def process_excel_file(self, file_path: str) -> Dict[str, Any]:
        """Traiter un fichier Excel et extraire les données"""
        if not EXCEL_AVAILABLE:
            return {'error': 'Support Excel non disponible. Installez openpyxl.'}
        
        try:
            # Lire toutes les feuilles
            excel_data = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
            
            analysis = {
                'type': 'excel',
                'sheet_count': len(excel_data),
                'sheets': {},
                'summary': {
                    'total_rows': 0,
                    'total_columns': 0,
                    'potential_financial_data': False
                }
            }
            
            financial_keywords = ['total', 'amount', 'price', 'cost', 'revenue', 'expense', 
                                'invoice', 'payment', 'balance', 'tax', 'vat', 'tva']
            
            for sheet_name, df in excel_data.items():
                sheet_info = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist(),
                    'sample_data': df.head(3).to_dict() if not df.empty else {},
                    'has_financial_indicators': False
                }
                
                # Vérifier si c'est des données financières
                column_text = ' '.join(str(col).lower() for col in df.columns)
                if any(keyword in column_text for keyword in financial_keywords):
                    sheet_info['has_financial_indicators'] = True
                    analysis['summary']['potential_financial_data'] = True
                
                # Détecter les colonnes numériques (potentielles données financières)
                numeric_columns = df.select_dtypes(include=['number']).columns.tolist()
                if numeric_columns:
                    sheet_info['numeric_columns'] = numeric_columns
                    sheet_info['numeric_summary'] = df[numeric_columns].describe().to_dict()
                
                analysis['sheets'][sheet_name] = sheet_info
                analysis['summary']['total_rows'] += len(df)
                analysis['summary']['total_columns'] += len(df.columns)
            
            return analysis
            
        except Exception as e:
            return {'error': f'Erreur lors du traitement Excel: {str(e)}'}
    
    def process_csv_file(self, file_path: str) -> Dict[str, Any]:
        """Traiter un fichier CSV"""
        try:
            # Essayer différents encodages et délimiteurs
            encodings = ['utf-8', 'latin-1', 'cp1252']
            separators = [',', ';', '\t']
            
            df = None
            used_encoding = None
            used_separator = None
            
            for encoding in encodings:
                for sep in separators:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding, sep=sep)
                        if len(df.columns) > 1:  # Bon délimiteur trouvé
                            used_encoding = encoding
                            used_separator = sep
                            break
                    except:
                        continue
                if df is not None:
                    break
            
            if df is None:
                return {'error': 'Impossible de lire le fichier CSV'}
            
            analysis = {
                'type': 'csv',
                'encoding': used_encoding,
                'separator': used_separator,
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist(),
                'sample_data': df.head(5).to_dict() if not df.empty else {},
                'data_types': df.dtypes.astype(str).to_dict(),
                'potential_financial_data': False
            }
            
            # Analyse des données financières
            financial_keywords = ['montant', 'prix', 'coût', 'total', 'tva', 'tax', 'invoice', 
                                'facture', 'payment', 'paiement', 'balance', 'solde']
            
            column_text = ' '.join(str(col).lower() for col in df.columns)
            if any(keyword in column_text for keyword in financial_keywords):
                analysis['potential_financial_data'] = True
            
            # Statistiques sur les colonnes numériques
            numeric_columns = df.select_dtypes(include=['number']).columns.tolist()
            if numeric_columns:
                analysis['numeric_columns'] = numeric_columns
                analysis['numeric_summary'] = df[numeric_columns].describe().to_dict()
            
            return analysis
            
        except Exception as e:
            return {'error': f'Erreur lors du traitement CSV: {str(e)}'}
    
    def process_pdf_file(self, file_path: str) -> Dict[str, Any]:
        """Traiter un fichier PDF et extraire le texte"""
        if not PDF_AVAILABLE:
            return {'error': 'Support PDF non disponible. Installez PyPDF2 et pdfplumber.'}
        
        try:
            # Utiliser pdfplumber pour une meilleure extraction
            import pdfplumber
            
            text_content = []
            tables_found = []
            page_count = 0
            
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extraire le texte
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append({
                            'page': page_num,
                            'text': page_text[:1000]  # Limiter pour éviter les gros volumes
                        })
                    
                    # Extraire les tableaux
                    tables = page.extract_tables()
                    if tables:
                        for table_num, table in enumerate(tables):
                            tables_found.append({
                                'page': page_num,
                                'table_num': table_num + 1,
                                'rows': len(table),
                                'columns': len(table[0]) if table else 0,
                                'sample_rows': table[:3]  # 3 premières lignes
                            })
            
            # Analyse du contenu
            full_text = ' '.join([page['text'] for page in text_content])
            financial_keywords = ['facture', 'invoice', 'total', 'montant', 'tva', 'tax', 
                                'payment', 'paiement', 'balance', 'solde', '€', '$']
            
            analysis = {
                'type': 'pdf',
                'page_count': page_count,
                'has_text': len(text_content) > 0,
                'has_tables': len(tables_found) > 0,
                'tables_count': len(tables_found),
                'text_pages': len(text_content),
                'potential_financial_document': any(keyword.lower() in full_text.lower() 
                                                  for keyword in financial_keywords),
                'text_sample': full_text[:500] if full_text else '',
                'tables_summary': tables_found
            }
            
            return analysis
            
        except Exception as e:
            return {'error': f'Erreur lors du traitement PDF: {str(e)}'}
    
    def process_image_file(self, file_path: str) -> Dict[str, Any]:
        """Traiter un fichier image avec OCR"""
        if not PIL_AVAILABLE:
            return {'error': 'Support image non disponible. Installez Pillow.'}
        
        try:
            with Image.open(file_path) as img:
                analysis = {
                    'type': 'image',
                    'format': img.format,
                    'mode': img.mode,
                    'size': img.size,
                    'width': img.size[0],
                    'height': img.size[1],
                    'has_text': False,
                    'extracted_text': '',
                    'potential_financial_document': False
                }
                
                # Tentative d'OCR si tesseract est disponible
                try:
                    import pytesseract
                    extracted_text = pytesseract.image_to_string(img, lang='eng+fra')
                    
                    if extracted_text.strip():
                        analysis['has_text'] = True
                        analysis['extracted_text'] = extracted_text[:1000]  # Limiter
                        
                        # Vérifier si c'est un document financier
                        financial_keywords = ['facture', 'invoice', 'total', 'montant', 'tva', 
                                            'tax', 'payment', '€', '$', 'balance', 'receipt']
                        
                        text_lower = extracted_text.lower()
                        if any(keyword in text_lower for keyword in financial_keywords):
                            analysis['potential_financial_document'] = True
                
                except ImportError:
                    analysis['ocr_available'] = False
                except Exception as ocr_error:
                    analysis['ocr_error'] = str(ocr_error)
                
                return analysis
                
        except Exception as e:
            return {'error': f'Erreur lors du traitement image: {str(e)}'}
    
    def process_word_file(self, file_path: str) -> Dict[str, Any]:
        """Traiter un fichier Word"""
        if not DOCX_AVAILABLE:
            return {'error': 'Support Word non disponible. Installez python-docx.'}
        
        try:
            doc = Document(file_path)
            
            # Extraire le texte
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extraire les tableaux
            tables_data = []
            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                tables_data.append({
                    'table_num': table_num,
                    'rows': len(table_data),
                    'columns': len(table_data[0]) if table_data else 0,
                    'sample_data': table_data[:3]
                })
            
            full_text = ' '.join(text_content)
            
            # Analyse financière
            financial_keywords = ['facture', 'invoice', 'total', 'montant', 'tva', 'tax',
                                'payment', 'paiement', 'balance', '€', '$']
            
            analysis = {
                'type': 'word',
                'paragraph_count': len(text_content),
                'table_count': len(tables_data),
                'has_tables': len(tables_data) > 0,
                'word_count': len(full_text.split()),
                'char_count': len(full_text),
                'text_sample': full_text[:500] if full_text else '',
                'tables_summary': tables_data,
                'potential_financial_document': any(keyword.lower() in full_text.lower() 
                                                  for keyword in financial_keywords)
            }
            
            return analysis
            
        except Exception as e:
            return {'error': f'Erreur lors du traitement Word: {str(e)}'}
    
    def process_text_file(self, file_path: str) -> Dict[str, Any]:
        """Traiter un fichier texte"""
        try:
            # Essayer différents encodages
            encodings = ['utf-8', 'latin-1', 'cp1252']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                return {'error': 'Impossible de lire le fichier texte'}
            
            lines = content.split('\n')
            
            analysis = {
                'type': 'text',
                'encoding': used_encoding,
                'line_count': len(lines),
                'word_count': len(content.split()),
                'char_count': len(content),
                'text_sample': content[:500] if content else '',
                'potential_financial_data': False
            }
            
            # Vérifier si c'est des données structurées (CSV-like)
            if ',' in content or ';' in content or '\t' in content:
                analysis['potential_structured_data'] = True
            
            # Analyse financière
            financial_keywords = ['facture', 'invoice', 'total', 'montant', 'tva', 'tax',
                                'payment', 'paiement', 'balance', '€', '$']
            
            content_lower = content.lower()
            if any(keyword in content_lower for keyword in financial_keywords):
                analysis['potential_financial_data'] = True
            
            return analysis
            
        except Exception as e:
            return {'error': f'Erreur lors du traitement texte: {str(e)}'}
    
    def process_json_file(self, file_path: str) -> Dict[str, Any]:
        """Traiter un fichier JSON"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            analysis = {
                'type': 'json',
                'structure': type(data).__name__,
                'sample_data': {},
                'potential_financial_data': False
            }
            
            if isinstance(data, dict):
                analysis['key_count'] = len(data)
                analysis['keys'] = list(data.keys())[:10]  # Premières 10 clés
                analysis['sample_data'] = {k: v for i, (k, v) in enumerate(data.items()) if i < 3}
            elif isinstance(data, list):
                analysis['item_count'] = len(data)
                if data and isinstance(data[0], dict):
                    analysis['sample_keys'] = list(data[0].keys())[:10]
                analysis['sample_data'] = data[:2] if len(data) > 0 else []
            
            # Analyse financière
            content_str = json.dumps(data).lower()
            financial_keywords = ['amount', 'price', 'cost', 'total', 'tax', 'invoice',
                                'payment', 'balance', 'montant', 'prix', 'tva']
            
            if any(keyword in content_str for keyword in financial_keywords):
                analysis['potential_financial_data'] = True
            
            return analysis
            
        except Exception as e:
            return {'error': f'Erreur lors du traitement JSON: {str(e)}'}
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Traiter un fichier selon son type"""
        try:
            file_info = self.get_file_info(file_path)
            extension = file_info['file_extension']
            
            # Traitement selon le type de fichier
            if extension in ['.xlsx', '.xls']:
                analysis = self.process_excel_file(file_path)
            elif extension == '.csv':
                analysis = self.process_csv_file(file_path)
            elif extension == '.pdf':
                analysis = self.process_pdf_file(file_path)
            elif extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']:
                analysis = self.process_image_file(file_path)
            elif extension in ['.docx', '.doc']:
                analysis = self.process_word_file(file_path)
            elif extension == '.txt':
                analysis = self.process_text_file(file_path)
            elif extension == '.json':
                analysis = self.process_json_file(file_path)
            else:
                analysis = {'error': f'Type de fichier non supporté: {extension}'}
            
            # Ajouter les informations de base
            if 'error' not in analysis:
                analysis.update(file_info)
                analysis['processing_timestamp'] = datetime.now().isoformat()
            
            return analysis
            
        except Exception as e:
            logger.error(f"Erreur lors du traitement du fichier {file_path}: {str(e)}")
            return {'error': f'Erreur générale: {str(e)}'}
    
    def save_uploaded_file(self, file, user_id: int, conversation_id: Optional[int] = None) -> Tuple[str, Dict[str, Any]]:
        """Sauvegarder un fichier uploadé et le traiter"""
        if not file or file.filename == '':
            raise ValueError("Aucun fichier fourni")
        
        if not self.is_allowed_file(file.filename):
            raise ValueError(f"Type de fichier non autorisé: {file.filename}")
        
        # Générer un nom de fichier sécurisé
        original_filename = file.filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        secure_name = secure_filename(original_filename)
        filename = f"{user_id}_{timestamp}_{secure_name}"
        
        # Créer le dossier utilisateur
        user_folder = os.path.join(self.upload_folder, str(user_id))
        os.makedirs(user_folder, exist_ok=True)
        
        file_path = os.path.join(user_folder, filename)
        
        # Sauvegarder le fichier
        file.save(file_path)
        
        # Vérifier la taille
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            os.remove(file_path)
            raise ValueError(f"Fichier trop volumineux: {file_size} bytes (max: {self.max_file_size})")
        
        # Traiter le fichier
        analysis = self.process_file(file_path)
        
        return file_path, analysis

# Instance globale du service
file_processor = FileProcessorService()